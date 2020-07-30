from docx import Document

document = Document("Forecast_Table_GSM.docx")

data = document.tables[0].rows[1:]
for i in data:
    predictor = "col"
    percent = int(i.cells[1].text[:-1])
    lag = int(i.cells[4].text)
    lag_type = "Use all lags"
    layer_num = int(i.cells[6].text)
    neuron_nums = [int(j) for j in i.cells[7].text.split(',')]
    activation = "Relu"
    epoch = int(i.cells[9].text)
    batch_size = int(i.cells[10].text)
    optimizer = "Adam"
    loss_function = i.cells[12].text.replace("\n", " ")
    learning_rate = 0.001
    pred_num = int(i.cells[14].text)
    print(percent, lag, lag_type, layer_num, neuron_nums, activation, epoch, batch_size, optimizer, loss_function, learning_rate, pred_num)

"""
table = document.add_table(rows=1, cols=17)
hdr_cells = table.rows[0].cells

hdr_cells[0].text = "Predictor-Target Seçimi"
hdr_cells[1].text = "Model Adı"
hdr_cells[2].text = "Eğitim Satır Yüzdesi/Satır Numarası"
hdr_cells[3].text = "Ölçek Türü"
hdr_cells[4].text = "Gecikme Sayısı"
hdr_cells[5].text = "Gecikme Seçeneği"
hdr_cells[6].text = "Gizli Katman Sayısı"
hdr_cells[7].text = "Nöron Sayısı"
hdr_cells[8].text = "Aktivasyon Fonsiyonu"
hdr_cells[9].text = "Epoch"
hdr_cells[10].text = "Batch Size"
hdr_cells[11].text = "Optimizer"
hdr_cells[12].text = "Kayıp Fonksiyonu"
hdr_cells[13].text = "Öğrenme Oranı"
hdr_cells[14].text = "Tahmin Sayısı"
hdr_cells[15].text = "MAPE"
"""
# document.save('demo.docx')

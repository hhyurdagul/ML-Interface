import tkinter as tk
from tkinter import ttk
from tkinter import filedialog
from pandastable import Table
from matplotlib import use as mat_backend
mat_backend("TkAgg")

import numpy as np
import pandas as pd
import matplotlib.pyplot as plt

from joblib import dump, load
from sklearn.svm import SVR, NuSVR
from sklearn.preprocessing import StandardScaler, MinMaxScaler
from sklearn.model_selection import GridSearchCV, train_test_split, cross_validate

import os
import json
from shutil import copyfile
from docx import Document
from pickle import dump as pickle_dump
from pickle import load as pickle_load

from .helpers import *

class SupportVectorMachine:
    def __init__(self):
        self.root = ttk.Frame()
        
        # Get Train Set
        get_train_set_frame = ttk.Labelframe(self.root, text="Get Train Set")
        get_train_set_frame.grid(column=0, row=0)

        file_path = tk.StringVar(value="")
        ttk.Label(get_train_set_frame, text="Train File Path").grid(column=0, row=0)
        ttk.Entry(get_train_set_frame, textvariable=file_path).grid(column=1, row=0)
        ttk.Button(get_train_set_frame, text="Read Data", command=lambda: self.readCsv(file_path)).grid(column=2, row=0)

        self.input_list = tk.Listbox(get_train_set_frame)
        self.input_list.grid(column=0, row=1)
        self.input_list.bind("<Double-Button-1>", self.addPredictor)
        self.input_list.bind("<Double-Button-3>", self.addTarget)

        self.predictor_list = tk.Listbox(get_train_set_frame)
        self.predictor_list.grid(column=1, row=1)
        self.predictor_list.bind("<Double-Button-1>", self.ejectPredictor)

        self.target_list = tk.Listbox(get_train_set_frame)
        self.target_list.grid(column=2, row=1)
        self.target_list.bind("<Double-Button-1>", self.ejectTarget)

        ttk.Button(get_train_set_frame, text="Add Predictor", command=self.addPredictor).grid(column=1, row=2)
        ttk.Button(get_train_set_frame, text="Eject Predictor", command=self.ejectPredictor).grid(column=1, row=3)

        ttk.Button(get_train_set_frame, text="Add Target", command=self.addTarget).grid(column=2, row=2)
        ttk.Button(get_train_set_frame, text="Eject Target", command=self.ejectTarget).grid(column=2, row=3)
       
        # Model testing and validation
        model_validation_frame = ttk.Labelframe(self.root, text="Model testing and validation")
        model_validation_frame.grid(column=0, row=1)

        self.do_forecast_option = tk.IntVar(value=0)
        tk.Checkbutton(model_validation_frame, text="Do Forecast", offvalue=0, onvalue=1, variable=self.do_forecast_option).grid(column=0, row=0, columnspan=2)
        
        self.validation_option = tk.IntVar(value=0)
        self.random_percent_var = tk.IntVar(value=70)
        self.cross_val_var = tk.IntVar(value=5)
        tk.Radiobutton(model_validation_frame, text="No validation, use all data rows", value=0, variable=self.validation_option).grid(column=0, row=1, columnspan=2, sticky=tk.W)
        tk.Radiobutton(model_validation_frame, text="Random percent", value=1, variable=self.validation_option).grid(column=0, row=2, sticky=tk.W)
        tk.Radiobutton(model_validation_frame, text="K-fold cross-validation", value=2, variable=self.validation_option).grid(column=0, row=3, sticky=tk.W)
        tk.Radiobutton(model_validation_frame, text="Leave one out cross-validation", value=3, variable=self.validation_option).grid(column=0, row=4, columnspan=2, sticky=tk.W)
        ttk.Entry(model_validation_frame, textvariable=self.random_percent_var, width=8).grid(column=1, row=2)
        ttk.Entry(model_validation_frame, textvariable=self.cross_val_var, width=8).grid(column=1, row=3)
        
        # Customize Train Set
        customize_train_set_frame = ttk.LabelFrame(self.root, text="Customize Train Set")
        customize_train_set_frame.grid(column=0, row=2)
        
        self.lookback_option = tk.IntVar(value=0)
        self.lookback_val_var = tk.IntVar(value="") # type: ignore
        tk.Checkbutton(customize_train_set_frame, text="Lookback", offvalue=0, onvalue=1, variable=self.lookback_option).grid(column=0, row=0)
        tk.Entry(customize_train_set_frame, textvariable=self.lookback_val_var, width=8).grid(column=1, row=0)

        self.seasonal_lookback_option = tk.IntVar(value=0)
        self.seasonal_period_var = tk.IntVar(value="") # type: ignore
        self.seasonal_val_var = tk.IntVar(value="") # type: ignore
        tk.Checkbutton(customize_train_set_frame, text="Periodic Lookback", offvalue=0, onvalue=1, variable=self.seasonal_lookback_option).grid(column=0, row=1)
        tk.Entry(customize_train_set_frame, textvariable=self.seasonal_period_var, width=9).grid(column=0, row=2)
        tk.Entry(customize_train_set_frame, textvariable=self.seasonal_val_var, width=8).grid(column=1, row=2)

        self.scale_var = tk.StringVar(value="None")
        ttk.Label(customize_train_set_frame, text="Scale Type").grid(column=0, row=3)
        ttk.OptionMenu(customize_train_set_frame, self.scale_var, "None", "None","StandardScaler", "MinMaxScaler").grid(column=1, row=3)

        # Model
        model_frame = ttk.Labelframe(self.root, text="Model Frame")
        model_frame.grid(column=1, row=0)

        ## Model Type
        model_type_frame = ttk.Labelframe(model_frame, text="Type of SVR Model")
        model_type_frame.grid(column=0, row=0)

        self.model_type_var = tk.IntVar(value=0)
        tk.Radiobutton(model_type_frame, text="Epsilon-SVR", value=0, variable=self.model_type_var, command=self.openEntries).grid(column=0, row=0)
        tk.Radiobutton(model_type_frame, text="Nu-SVR", value=1, variable=self.model_type_var, command=self.openEntries).grid(column=1, row=0)

        ## Kernel Type
        kernel_type_frame = ttk.Labelframe(model_frame, text="Kernel Function")
        kernel_type_frame.grid(column=0, row=1)

        self.kernel_type_var = tk.IntVar()
        tk.Radiobutton(kernel_type_frame, text="Linear", value=0, variable=self.kernel_type_var, command=self.openEntries).grid(column=0, row=0, sticky=tk.W)
        tk.Radiobutton(kernel_type_frame, text="RBF", value=1, variable=self.kernel_type_var, command=self.openEntries).grid(column=1, row=0, sticky=tk.W)
        tk.Radiobutton(kernel_type_frame, text="Polynomial", value=2, variable=self.kernel_type_var, command=self.openEntries).grid(column=0, row=1)
        tk.Radiobutton(kernel_type_frame, text="Sigmoid", value=3, variable=self.kernel_type_var, command=self.openEntries).grid(column=1, row=1)
        
        ## Parameter Optimization
        parameter_optimization_frame = ttk.Labelframe(model_frame, text="Parameter Optimization")
        parameter_optimization_frame.grid(column=0, row=2)

        self.grid_option_var = tk.IntVar(value=0)
        tk.Checkbutton(parameter_optimization_frame, text="Do grid search for optimal parameters", offvalue=0, onvalue=1, variable=self.grid_option_var, command=self.openEntries).grid(column=0, row=0, columnspan=3)

        self.interval_var = tk.IntVar(value="") # type: ignore
        ttk.Label(parameter_optimization_frame, text="Interval:").grid(column=0, row=1)
        self.interval_entry = ttk.Entry(parameter_optimization_frame, textvariable=self.interval_var, width=8, state=tk.DISABLED)
        self.interval_entry.grid(column=1, row=1, pady=2)

        self.gs_cross_val_option = tk.IntVar(value=0)
        self.gs_cross_val_var = tk.IntVar(value="") # type: ignore
        tk.Checkbutton(parameter_optimization_frame, text="Cross validate; folds:", offvalue=0, onvalue=1, variable=self.gs_cross_val_option, command=self.openEntries).grid(column=0, row=2)
        self.gs_cross_val_entry = tk.Entry(parameter_optimization_frame, textvariable=self.gs_cross_val_var, state=tk.DISABLED, width=8)
        self.gs_cross_val_entry.grid(column=1, row=2)

        ## Model Parameters
        model_parameters_frame = ttk.LabelFrame(model_frame, text="Model Parameters")
        model_parameters_frame.grid(column=1, row=0, rowspan=3, columnspan=2)
        
        parameter_names = ["Epsilon", "Nu", "C", "Gamma", "Coef0", "Degree"]
        self.parameters = [tk.Variable(value="0.1"), tk.Variable(value="0.5"), tk.Variable(value="1"), tk.Variable(value="1"), tk.Variable(value="0"), tk.Variable(value="3")]
        self.optimization_parameters = [[tk.Variable(), tk.Variable()], [tk.Variable(), tk.Variable()], [tk.Variable(), tk.Variable()], [tk.Variable(), tk.Variable()], [tk.Variable(), tk.Variable()], [tk.Variable(), tk.Variable()]]
        
        ttk.Label(model_parameters_frame, text="Current").grid(column=1, row=0)
        ttk.Label(model_parameters_frame, text="------ Search Range ------").grid(column=2, row=0, columnspan=2)

        self.model_parameters_frame_options = [
            [
                ttk.Label(model_parameters_frame, text=j+":").grid(column=0, row=i+1),
                ttk.Entry(model_parameters_frame, textvariable=self.parameters[i], state=tk.DISABLED, width=9),
                ttk.Entry(model_parameters_frame, textvariable=self.optimization_parameters[i][0], state=tk.DISABLED, width=9),
                ttk.Entry(model_parameters_frame, textvariable=self.optimization_parameters[i][1], state=tk.DISABLED, width=9)
            ] for i,j in enumerate(parameter_names)
        ]

        for i, j in enumerate(self.model_parameters_frame_options):
            j[1].grid(column=1, row=i+1, padx=2, pady=2)
            j[2].grid(column=2, row=i+1, padx=2, pady=2)
            j[3].grid(column=3, row=i+1, padx=2, pady=2)

        ttk.Label(model_parameters_frame, text="Gamma:").grid(column=0, row=7)
        self.gamma_choice = tk.IntVar(value=0)
        self.gamma_radios = [tk.Radiobutton(model_parameters_frame, text=j, value=i, variable=self.gamma_choice, state=tk.DISABLED, command=self.openEntries) for i, j in enumerate(["Scale", "Auto", "Value"])]
        
        for i, j in enumerate(self.gamma_radios):
            j.grid(column=i+1, row=7)
    
        ttk.Button(model_frame, text="Create Model", command=self.createModel).grid(column=0, row=3)

        ttk.Button(model_frame, text="Save Model", command=self.saveModel).grid(column=1, row=3)
        ttk.Button(model_frame, text="Load Model", command=self.loadModel).grid(column=2, row=3)

        # Test Model
        test_model_frame = ttk.LabelFrame(self.root, text="Test Frame")
        test_model_frame.grid(column=1, row=1)

        ## Test Model Main
        test_model_main_frame = ttk.LabelFrame(test_model_frame, text="Test Model")
        test_model_main_frame.grid(column=0, row=0)

        forecast_num = tk.IntVar(value="") # type: ignore
        ttk.Label(test_model_main_frame, text="# of Forecast").grid(column=0, row=0)
        ttk.Entry(test_model_main_frame, textvariable=forecast_num).grid(column=1, row=0)
        ttk.Button(test_model_main_frame, text="Values", command=self.showPredicts).grid(column=2, row=0)

        test_file_path = tk.StringVar()
        ttk.Label(test_model_main_frame, text="Test File Path").grid(column=0, row=1)
        ttk.Entry(test_model_main_frame, textvariable=test_file_path).grid(column=1, row=1)
        ttk.Button(test_model_main_frame, text="Get Test Set", command=lambda: self.getTestSet(test_file_path)).grid(column=2, row=1)

        ttk.Button(test_model_main_frame, text="Test Model", command=lambda: self.forecast(forecast_num.get())).grid(column=2, row=3)
        ttk.Button(test_model_main_frame, text="Actual vs Forecast Graph", command=self.vsGraph).grid(column=0, row=4, columnspan=3)

        ## Test Model Metrics
        test_model_metrics_frame = ttk.LabelFrame(test_model_frame, text="Test Metrics")
        test_model_metrics_frame.grid(column=1, row=0)

        test_metrics = ["NMSE", "RMSE", "MAE", "MAPE", "SMAPE"]
        self.test_metrics_vars = [tk.Variable(), tk.Variable(), tk.Variable(), tk.Variable(), tk.Variable(), tk.Variable()]
        for i, j in enumerate(test_metrics):
            ttk.Label(test_model_metrics_frame, text=j).grid(column=0, row=i)
            ttk.Entry(test_model_metrics_frame, textvariable=self.test_metrics_vars[i]).grid(column=1,row=i)
        
        table_frame = ttk.LabelFrame(self.root, text="Table Operations")
        table_frame.grid(column=1, row=2)

        ttk.Button(table_frame, text="Create Table", command=self.createTable).grid(column=0, row=0)
        ttk.Button(table_frame, text="Add to Table", command=self.addToTable).grid(column=1, row=0)

        self.openEntries()


    def readCsv(self, file_path):
        path = filedialog.askopenfilename(filetypes=[("Csv Files", "*.csv"), ("Excel Files", "*.xl*")])
        file_path.set(path)
        if path.endswith(".csv"):
            self.df = pd.read_csv(path) # type: ignore
        else:
            try:
                self.df = pd.read_excel(path)
            except:
                self.df = pd.read_excel(path, engine="openpyxl")
        self.fillInputList()
        
    def fillInputList(self):
        self.input_list.delete(0, tk.END)

        self.df: pd.DataFrame
        for i in self.df.columns.to_list():
            self.input_list.insert(tk.END, i)

    def getTestSet(self, file_path):
        path = filedialog.askopenfilename(filetypes=[("Csv Files", "*.csv"), ("Excel Files", "*.xl*")])
        file_path.set(path)
        if path.endswith(".csv"):
            self.test_df = pd.read_csv(path) # type: ignore
        else:
            try:
                self.test_df = pd.read_excel(path)
            except:
                self.test_df = pd.read_excel(path, engine="openpyxl")

    def showPredicts(self):
        top = tk.Toplevel(self.root)
        df = pd.DataFrame({"Test": self.y_test, "Predict": self.pred})
        pt = Table(top, dataframe=df, editable=False)
        pt.show()

    def addPredictor(self, event=None):
        try:
            a = self.input_list.get(self.input_list.curselection())
            if a not in self.predictor_list.get(0,tk.END):
                self.predictor_list.insert(tk.END, a)
        except:
            pass

    def ejectPredictor(self, event=None):
        try:
            self.predictor_list.delete(self.predictor_list.curselection())
        except:
            pass
    
    def addTarget(self, event=None):
        try:
            a = self.input_list.get(self.input_list.curselection())
            if self.target_list.size() < 1:
                self.target_list.insert(tk.END, a)
        except:
            pass

    def ejectTarget(self, event=None):
        try:
            self.target_list.delete(self.target_list.curselection())
        except:
            pass

    def createTable(self):
        old = os.path.join(os.path.abspath("."), "Table Templates/SVM.docx")
        path = filedialog.asksaveasfilename()+".docx"
        copyfile(old, path)
        document = Document(path)
        self.document_path = path
        self.document = document

    def addToTable(self):
        table = self.document.tables[0]
        table_style = table.style
        style = table.cell(0, 0).paragraphs[-1].style

        cells = table.add_row().cells
        cells[0].text = "%"+str(self.random_percent_var.get()) if self.validation_option.get() == 1 else "%100"
        cells[1].text = str(self.lookback_val_var.get()) if self.lookback_option.get() else "-"
        cells[2].text = str(self.seasonal_period_var.get())+ "-" + str(self.seasonal_val_var.get()) if self.seasonal_lookback_option.get() else "-"
        cells[3].text = "Nu" if self.model_type_var.get() else "Epsilon"
        cells[4].text = "Min-Max" if self.scale_var.get() == "MinMaxScaler" else "Standart" if self.scale_var.get() == "StandardScaler" else "Yok"
        cells[5].text = ["Doğrusal", "RBF", "Polinom", "Sigmoid"][self.kernel_type_var.get()]
        t_ep = "{:.2e}".format(float(self.parameters[0].get()))
        cells[6].text = t_ep[0] + "*10"
        c = cells[6].paragraphs[0]
        s = c.add_run(t_ep[-1])
        s.font.superscript = True
        cells[7].text = "2"
        c = cells[7].paragraphs[0]
        s = c.add_run(str(float(self.parameters[2].get()))[:-2])
        s.font.superscript = True
        if self.gamma_choice.get() == 2:
            cells[8].text = "2"
            c = cells[8].paragraphs[0]
            s = c.add_run(str(self.parameters[3].get()))
            s.font.superscript = True
        else:
            cells[8].text = "-"
        if self.kernel_type_var.get() == 2:
            cells[9].text = self.parameters[-1].get()
        else:
            cells[9].text = "-"
        cells[10].text = "Değer" if self.gamma_choice.get() == 2 else "Otomatik" if self.gamma_choice.get() == 1 else "Ölçeklendirilmiş"
        cells[11].text = self.test_metrics_vars[2].get()

        for i in cells:
            for j in i.paragraphs:
                j.style = style
        table.style = table_style
        self.document.save(self.document_path)
    
    def saveModel(self):
        path = filedialog.asksaveasfilename()
        params = self.model.get_params()
        params["predictor_names"] = self.predictor_names
        params["label_name"] = self.label_name
        params["is_round"] = self.is_round
        params["is_negative"] = self.is_negative
        params["do_forecast"] = self.do_forecast_option.get()
        params["validation_option"] = self.validation_option.get()
        params["random_percent"] = self.random_percent_var.get() if self.validation_option.get() == 1 else None
        params["k_fold_cv"] = self.cross_val_var.get() if self.validation_option.get() == 2 else None
        params["lookback_option"] = self.lookback_option.get()
        params["lookback_value"] = self.lookback_val_var.get() if self.lookback_option.get() else None
        params["seasonal_lookback_option"] = self.seasonal_lookback_option.get()
        params["seasonal_period"] = self.seasonal_period_var.get() if self.seasonal_lookback_option.get() else None
        params["seasonal_value"] = self.seasonal_val_var.get() if self.seasonal_lookback_option.get() else None
        params["sliding"] = self.sliding
        params["scale_type"] = self.scale_var.get()

        os.mkdir(path)
        dump(self.model, path+"/model.joblib")
        if self.scale_var.get() != "None":
            with open(path+"/feature_scaler.pkl", "wb") as f:
                pickle_dump(self.feature_scaler, f)
            with open(path+"/label_scaler.pkl", "wb") as f:
                pickle_dump(self.label_scaler, f)
        if self.lookback_option.get() == 1:
            with open(path+"/last_values.npy", 'wb') as outfile:
                np.save(outfile, self.last)
        if self.seasonal_lookback_option.get() == 1:
            with open(path+"/seasonal_last_values.npy", 'wb') as outfile:
                np.save(outfile, self.seasonal_last)
        with open(path+"/model.json", 'w') as outfile:
            json.dump(params, outfile)

    def loadModel(self):
        path = filedialog.askdirectory()
        model_path = path + "/model.joblib"
        self.model = load(model_path)
        infile = open(path+"/model.json")
        params = json.load(infile)

        self.predictor_names = params["predictor_names"]
        self.label_name = params["label_name"]
        try:
            self.is_round = params["is_round"]
        except:
            self.is_round = True
        try:
            self.is_negative = params["is_negative"]
        except:
            self.is_negative = False
        self.do_forecast_option.set(params["do_forecast"])
        self.validation_option.set(params["validation_option"])
        if params["validation_option"] == 1:
            self.random_percent_var.set(params["random_percent"])
        elif params["validation_option"] == 2:
            self.cross_val_var.set(params["k_fold_cv"])
        self.lookback_option.set(params["lookback_option"]) 
        self.sliding=-1
        if params["lookback_option"] == 1:
            self.lookback_val_var.set(params["lookback_value"])
            last_values = open(path+"/last_values.npy", 'rb')
            self.last = np.load(last_values)
            last_values.close()
        try:
            self.sliding = params["sliding"]
            self.seasonal_lookback_option.set(params["seasonal_lookback_option"]) 
            if params["seasonal_lookback_option"] == 1:
                self.seasonal_period_var.set(params["seasonal_period"])
                self.seasonal_val_var.set(params["seasonal_value"])
                seasonal_last_values = open(path+"/seasonal_last_values.npy", 'rb')
                self.seasonal_last = np.load(seasonal_last_values)
                seasonal_last_values.close()
        except:
            pass
        self.scale_var.set(params["scale_type"])
        if params["scale_type"] != "None":
            try:
                with open(path+"/feature_scaler.pkl", "rb") as f:
                    self.feature_scaler = pickle_load(f)
                with open(path+"/label_scaler.pkl", "rb") as f:
                    self.label_scaler = pickle_load(f)
            except:
                pass
        try:
            self.parameters[0].set(params["epsilon"])
            self.model_type_var.set(0)
        except:
            self.parameters[1].set(params["nu"])
            self.model_type_var.set(1)
        self.parameters[2].set(np.log2(params["C"]))
        if params["gamma"] == "scale":
            self.gamma_choice.set(0)
        elif params["gamma"] == "auto":
            self.gamma_choice.set(1)
        else:
            self.gamma_choice.set(2)
            self.parameters[3].set(np.log2(params["gamma"]))
        self.parameters[4].set(params["coef0"])
        self.parameters[5].set(params["degree"])
        kernel = 0 if params["kernel"] == "linear" else 1 if params["kernel"] == "rbf" else 2 if params["kernel"] == "poly" else 3
        self.kernel_type_var.set(kernel)
       
        self.openEntries()
        msg = f"Predictor names are {self.predictor_names}\nLabel name is {self.label_name}"
        popupmsg(msg)

    def openEntries(self):
        to_open = []
        for i in self.model_parameters_frame_options:
            i[1]["state"] = tk.DISABLED
            i[2]["state"] = tk.DISABLED
            i[3]["state"] = tk.DISABLED

        self.interval_entry["state"] = tk.DISABLED
        self.gs_cross_val_entry["state"] = tk.DISABLED

        if self.kernel_type_var.get() == 0:
            for i in self.gamma_radios:
                i["state"] = tk.DISABLED
        else:
            for i in self.gamma_radios:
                i["state"] = tk.NORMAL
            if self.gamma_choice.get() == 2:
                to_open.append(3)
        
        if self.gs_cross_val_option.get() == 1:
            self.gs_cross_val_entry["state"] = tk.NORMAL

        if self.kernel_type_var.get() == 2:
            to_open.append(5)
            to_open.append(4)
        elif self.kernel_type_var.get() == 3:
            to_open.append(4) 

        to_open.append(self.model_type_var.get())
        to_open.append(2)

        to_open.sort()
        
        opt = self.grid_option_var.get()

        self.open(to_open, opt)

    def open(self, to_open, opt=0):
        if opt == 1:
            self.interval_entry["state"] = tk.NORMAL
            for i in to_open:
                self.model_parameters_frame_options[i][2]["state"] = tk.NORMAL
                self.model_parameters_frame_options[i][3]["state"] = tk.NORMAL
        else:
            for i in to_open:
                self.model_parameters_frame_options[i][1]["state"] = tk.NORMAL
        
        self.vars_nums = to_open

    def getLookback(self, X, y, lookback=0, seasons=0, seasonal_lookback=0, sliding=-1):
        if sliding == 0:
            for i in range(1, lookback+1):
                X[f"t-{i}"] = y.shift(i)
        elif sliding == 1:
            for i in range(1, seasons+1):
                X[f"t-{i*seasonal_lookback}"] = y.shift(i*seasonal_lookback)
        elif sliding == 2:
            for i in range(1, lookback+1):
                X[f"t-{i}"] = y.shift(i)
            for i in range(1, seasons+1):
                X[f"t-{i*seasonal_lookback}"] = y.shift(i*seasonal_lookback)

        X.dropna(inplace=True)
        a = X.to_numpy()
        b = y.iloc[-len(a):].to_numpy().reshape(-1)
        
        if sliding == 0:
            self.last = b[-lookback:]
        elif sliding == 1:
            self.seasonal_last = b[-seasonal_lookback*seasons:]
        elif sliding == 2:
            self.last = b[-(lookback+seasonal_lookback):-seasonal_lookback]
            self.seasonal_last = b[-seasonal_lookback*seasons:]

        return a, b

    def getData(self):
        self.is_round = False
        self.is_negative = False
        lookback_option = self.lookback_option.get()
        seasonal_lookback_option = self.seasonal_lookback_option.get()
        sliding = lookback_option + 2*seasonal_lookback_option - 1
        self.sliding = sliding
        scale_choice = self.scale_var.get()

        self.predictor_names = list(self.predictor_list.get(0, tk.END))
        self.label_name = self.target_list.get(0)

        self.df: pd.DataFrame
        X = self.df[self.predictor_names].copy()
        y = self.df[self.label_name].copy()
        
        if y.dtype == int or y.dtype == np.int or y.dtype == np.int64:
            self.is_round = True
        if any(y < 0):
            self.is_negative = True

        if scale_choice == "StandardScaler":
            self.feature_scaler = StandardScaler()
            self.label_scaler = StandardScaler()
        
            X.iloc[:] = self.feature_scaler.fit_transform(X)
            y.iloc[:] = self.label_scaler.fit_transform(y.values.reshape(-1,1)).reshape(-1)
        
        elif scale_choice == "MinMaxScaler":
            self.feature_scaler = MinMaxScaler()
            self.label_scaler = MinMaxScaler()
            
            X.iloc[:] = self.feature_scaler.fit_transform(X)
            y.iloc[:] = self.label_scaler.fit_transform(y.values.reshape(-1,1)).reshape(-1)
        
        try:
            lookback = self.lookback_val_var.get()
        except:
            lookback = 0
        try:
            seasonal_period = self.seasonal_period_var.get()
            seasonal_lookback = self.seasonal_val_var.get()
        except:
            seasonal_period = 0
            seasonal_lookback = 0
            
        X,y = self.getLookback(X, y, lookback, seasonal_period, seasonal_lookback, sliding)

        return X, y


    def createModel(self):
        gamma_choice = self.gamma_choice.get()
        kernels = ["linear", "rbf", "poly", "sigmoid"]
        kernel = kernels[self.kernel_type_var.get()]
        
        do_forecast = self.do_forecast_option.get()
        val_option = self.validation_option.get()
        
        X, y = self.getData()
        X: np.ndarray
        y: np.ndarray

        if self.grid_option_var.get() == 0:
            epsilon = float(self.parameters[0].get())
            nu = float(self.parameters[1].get())
            C = 2 ** float(self.parameters[2].get())
            gamma = 2 ** float(self.parameters[3].get()) if gamma_choice == 2 else "auto" if gamma_choice == 1 else "scale"
            coef0 = float(self.parameters[4].get())
            degree = float(self.parameters[5].get())
            
            if self.model_type_var.get() == 0:
                model = SVR(kernel=kernel, C=C, epsilon=epsilon, gamma=gamma, coef0=coef0, degree=degree)
            else:
                model = NuSVR(kernel=kernel, C=C, nu=nu, gamma=gamma, coef0=coef0, degree=degree)

            if val_option == 0:
                model.fit(X, y)
                if do_forecast == 0:
                    pred = model.predict(X).reshape(-1)
                    if self.scale_var.get() != "None":
                        pred = self.label_scaler.inverse_transform(pred.reshape(-1,1)).reshape(-1) # type: ignore
                        y = self.label_scaler.inverse_transform(y.reshape(-1,1)).reshape(-1) # type: ignore
                    losses = loss(y, pred)[:-1]
                    self.y_test = y
                    self.pred = pred
                    for i,j in enumerate(losses):
                        self.test_metrics_vars[i].set(j)
                self.model = model # type: ignore
            
            elif val_option == 1:
                if do_forecast == 0:
                    X_train, X_test, y_train, y_test = train_test_split(X,y, train_size=self.random_percent_var.get()/100)
                    model.fit(X_train, y_train)
                    pred = model.predict(X_test).reshape(-1)
                    if self.scale_var.get() != "None":
                        pred = self.label_scaler.inverse_transform(pred.reshape(-1,1)).reshape(-1) # type: ignore
                        y_test = self.label_scaler.inverse_transform(y_test.reshape(-1,1)).reshape(-1) # type: ignore
                    losses = loss(y_test, pred)[:-1]
                    self.y_test = y_test
                    self.pred = pred
                    for i,j in enumerate(losses):
                        self.test_metrics_vars[i].set(j)
                else:
                    size = int((self.random_percent_var.get()/100)*len(X))
                    X = X[-size:]
                    y = y[-size:]
                    model.fit(X, y)
                self.model = model # type: ignore

            elif val_option == 2:
                if do_forecast == 0:
                    cvs = cross_validate(model, X, y, cv=self.cross_val_var.get(), scoring=skloss)
                    for i,j in enumerate(list(cvs.values())[2:]):
                        self.test_metrics_vars[i].set(j.mean())

            elif val_option == 3:
                if do_forecast == 0:
                    cvs = cross_validate(model, X, y, cv=X.shape[0]-1, scoring=skloss)
                    for i,j in enumerate(list(cvs.values())[2:]):
                        self.test_metrics_vars[i].set(j.mean())
            
        else:
            params = {}
            interval = self.interval_var.get()
             
            params["C"] = np.unique(np.logspace(float(self.optimization_parameters[2][0].get()), float(self.optimization_parameters[2][1].get()), interval, base=2))
            if self.model_type_var.get() == 0:
                params["epsilon"] = np.unique(np.linspace(float(self.optimization_parameters[0][0].get()), float(self.optimization_parameters[0][1].get()), interval))
                model = SVR()
            else:
                min_nu = max(0.0001, float(self.optimization_parameters[1][0].get()))
                max_nu = min(1, float(self.optimization_parameters[1][1].get()))
                params["nu"] = np.unique(np.linspace(min_nu, max_nu, interval))
                model = NuSVR()
            if kernel != "linear":
                if gamma_choice == 2:
                    params["gamma"] = np.unique(np.logspace(float(self.optimization_parameters[3][0].get()), float(self.optimization_parameters[3][1].get()), interval, base=2))
                elif gamma_choice == 1:
                    params["gamma"] = ["auto"]
                else:
                    params["gamma"] = ["scale"]
            
            if kernel == "poly" or kernel == "sigmoid":
                params["coef0"] = np.unique(np.linspace(float(self.optimization_parameters[4][0].get()), float(self.optimization_parameters[4][1].get()), interval))

            if kernel == "poly":
                params["degree"] = np.unique(np.linspace(float(self.optimization_parameters[5][0].get()), float(self.optimization_parameters[5][1].get()), interval, dtype=int))

            params["kernel"] = [kernel]

            cv = self.gs_cross_val_var.get() if self.gs_cross_val_option.get() == 1 else None
            
            regressor = GridSearchCV(model, params, cv=cv)
            
            if val_option == 0:
                regressor.fit(X, y)
                if do_forecast == 0:
                    pred = regressor.predict(X)
                    if self.scale_var.get() != "None":
                        pred = self.label_scaler.inverse_transform(pred.reshape(-1,1)).reshape(-1) # type: ignore
                        y = self.label_scaler.inverse_transform(y.reshape(-1,1)).reshape(-1) # type: ignore
                    losses = loss(y, pred)[:-1]
                    self.y_test = y
                    self.pred = pred
                    for i,j in enumerate(losses):
                        self.test_metrics_vars[i].set(j)
                self.model = regressor.best_estimator_

            elif val_option == 1:
                if do_forecast == 0:
                    X_train, X_test, y_train, y_test = train_test_split(X,y, train_size=self.random_percent_var.get()/100)
                    regressor.fit(X_train, y_train)
                    pred = regressor.predict(X_test)
                    if self.scale_var.get() != "None":
                        pred = self.label_scaler.inverse_transform(pred.reshape(-1,1)).reshape(-1) # type: ignore
                        y_test = self.label_scaler.inverse_transform(y_test.reshape(-1,1)).reshape(-1) # type: ignore
                    losses = loss(y_test, pred)[:-1]
                    self.y_test = y_test
                    self.pred = pred
                    for i,j in enumerate(losses):
                        self.test_metrics_vars[i].set(j)
                else:
                    size = int((self.random_percent_var.get()/100)*len(X))
                    X = X[-size:]
                    y = y[-size:]
                    regressor.fit(X, y)
                self.model = regressor.best_estimator_
            
            popupmsg("Best Params: " + str(self.model.get_params()))
        
    def forecastLookback(self, num, lookback=0, seasons=0, seasonal_lookback=0, sliding=-1):
        self.test_df: pd.DataFrame
        pred = []
        if sliding == 0:
            last = self.last
            for i in range(num):
                X_test = self.test_df[self.predictor_names].iloc[i]
                if self.scale_var.get() != "None":
                    X_test.iloc[:] = self.feature_scaler.transform(X_test.values.reshape(1,-1)).reshape(-1) # type: ignore
                for j in range(1, lookback+1):
                    X_test[f"t-{j}"] = last[-j] # type: ignore
                to_pred = X_test.to_numpy().reshape(1,-1) # type: ignore
                out = self.model.predict(to_pred)
                last = np.append(last, out)[-lookback:]
                pred.append(out)

        elif sliding == 1:
            seasonal_last = self.seasonal_last
            for i in range(num):
                X_test = self.test_df[self.predictor_names].iloc[i]
                if self.scale_var.get() != "None":
                    X_test.iloc[:] = self.feature_scaler.transform(X_test.values.reshape(1,-1)).reshape(-1) # type: ignore
                for j in range(1, seasons+1):
                    X_test[f"t-{j*seasonal_last}"] = seasonal_last[-j*seasonal_lookback] # type: ignore
                to_pred = X_test.to_numpy().reshape(1,-1) # type: ignore
                out = self.model.predict(to_pred)
                seasonal_last = np.append(seasonal_last, out)[1:]
                pred.append(out)

        elif sliding == 2:
            last = self.last
            seasonal_last = self.seasonal_last
            for i in range(num):
                X_test = self.test_df[self.predictor_names].iloc[i]
                if self.scale_var.get() != "None":
                    X_test.iloc[:] = self.feature_scaler.transform(X_test.values.reshape(1,-1)).reshape(-1) # type: ignore
                for j in range(1, lookback+1):
                    X_test[f"t-{j}"] = last[-j] # type: ignore
                for j in range(1, seasons+1):
                    X_test[f"t-{j*seasonal_lookback}"] = seasonal_last[-j*seasonal_lookback] # type: ignore
                to_pred = X_test.to_numpy().reshape(1,-1) # type: ignore
                out = self.model.predict(to_pred)
                last = np.append(last, out)[-lookback:]
                seasonal_last = np.append(seasonal_last, out)[1:]
                pred.append(out)

        return np.array(pred).reshape(-1)

    def forecast(self, num):
        lookback_option = self.lookback_option.get()
        seasonal_lookback_option = self.seasonal_lookback_option.get()
        X_test = self.test_df[self.predictor_names][:num].to_numpy() # type: ignore
        y_test = self.test_df[self.label_name][:num].to_numpy().reshape(-1) # type: ignore
        self.y_test = y_test
       
        if lookback_option == 0 and seasonal_lookback_option == 0:
            if self.scale_var.get() != "None":
                X_test = self.feature_scaler.transform(X_test)
            self.pred = self.model.predict(X_test).reshape(-1)
        else:
            sliding = self.sliding
            try:
                lookback = self.lookback_val_var.get()
            except:
                lookback = 0
            try:
                seasonal_lookback = self.seasonal_val_var.get()
                seasons = self.seasonal_period_var.get()
            except:
                seasonal_lookback = 0
                seasons = 0 

            self.pred = self.forecastLookback(num, lookback, seasons, seasonal_lookback, sliding)

        if self.scale_var.get() != "None":
            self.pred = self.label_scaler.inverse_transform(self.pred.reshape(-1,1)).reshape(-1) # type: ignore

        if not self.is_negative:
            self.pred = self.pred.clip(0, None)
        if self.is_round:
            self.pred = np.round(self.pred).astype(int)
        
        losses = loss(y_test, self.pred)
        for i in range(6):
            self.test_metrics_vars[i].set(losses[i])

    def vsGraph(self):
        y_test = self.y_test
        pred = self.pred
        plt.plot(y_test)
        plt.plot(pred)
        plt.legend(["test", "pred"], loc="upper left")
        plt.show()

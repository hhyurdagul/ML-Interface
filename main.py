# Gui
import tkinter as tk
from tkinter import ttk
from tkinter import filedialog
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg, NavigationToolbar2Tk
from pandastable import Table

# Self Use
from docx import Document

# Data
import pandas as pd
import numpy as np
from statsmodels.graphics.tsaplots import plot_acf, plot_pacf
from statsmodels.tsa.stattools import acf
from datetime import datetime

# Keras
from tensorflow.keras.backend import clear_session
from tensorflow.keras.models import Sequential
from tensorflow.keras.layers import Conv1D, MaxPooling1D
from tensorflow.keras.layers import Input, Flatten, Dense, LSTM, Bidirectional
from tensorflow.keras.optimizers import Adam, SGD, RMSprop
#from keras.models import Sequential
#from keras.layers import Input, Flatten, Dense, LSTM, Bidirectional
#from keras.optimizers import Adam, SGD, RMSprop
from kerastuner.tuners import RandomSearch

# Sklearn
from sklearn.svm import SVR, NuSVR
from sklearn.preprocessing import StandardScaler, MinMaxScaler

document = Document("Real Data/Load.docx")
rows = document.tables[0].rows[1:]

savedoc = Document("Real Data/Forecast_Bidirectional LSTM.docx")
savetable = savedoc.tables[0]

class GUI:
    def __init__(self):
        self.gui = tk.Tk()
        self.parent = ttk.Notebook(self.gui)

        time_series = TimeSeries()
        self.add(time_series, "Time Series")
        
        #svm = SupportVectorMachine()
        #self.add(svm, "Support Vector Machine")
      
        self.parent.pack(expand=1, fill='both')

    def add(self, frame, text):
        self.parent.add(frame.root, text=text)

    def start(self):
        self.gui.mainloop()


class TimeSeries:
    def __init__(self):
        self.root = ttk.Frame()
        
        # Get Train Set
        get_train_set_frame = ttk.Labelframe(self.root, text="Get Train Set")
        get_train_set_frame.grid(column=0, row=0)

        file_path = tk.StringVar(value="")
        ttk.Label(get_train_set_frame, text="Train File Path").grid(column=0, row=0)
        ttk.Entry(get_train_set_frame, textvariable=file_path).grid(column=1, row=0)
        ttk.Button(get_train_set_frame, text="Read Csv", command=lambda: self.readCsv(file_path)).grid(column=2, row=0)

        self.input_list = tk.Listbox(get_train_set_frame)
        self.input_list.grid(column=0, row=1)

        self.predictor_list = tk.Listbox(get_train_set_frame)
        self.predictor_list.grid(column=1, row=1)

        self.target_list = tk.Listbox(get_train_set_frame)
        self.target_list.grid(column=2, row=1)

        ttk.Button(get_train_set_frame, text="Add Predictor", command=self.addPredictor).grid(column=1, row=2)
        ttk.Button(get_train_set_frame, text="Eject Predictor", command=self.ejectPredictor).grid(column=1, row=3)

        ttk.Button(get_train_set_frame, text="Add Target", command=self.addTarget).grid(column=2, row=2)
        ttk.Button(get_train_set_frame, text="Eject Target", command=self.ejectTarget).grid(column=2, row=3)
       
        # Customize Train Set
        customize_train_set_frame = ttk.Labelframe(self.root, text="Customize Train Set")
        customize_train_set_frame.grid(column=0, row=1)

        self.train_size_var = tk.IntVar(value="")
        ttk.Label(customize_train_set_frame, text="# of Rows in Train Set").grid(column=0, row=0)
        ttk.Entry(customize_train_set_frame, textvariable=self.train_size_var).grid(column=1, row=0)

        self.size_choice_var = tk.IntVar(value=0)
        tk.Radiobutton(customize_train_set_frame, text="As Percent", value=0, variable=self.size_choice_var).grid(column=0, row=1)
        tk.Radiobutton(customize_train_set_frame, text="As Number", value=1, variable=self.size_choice_var).grid(column=1, row=1)

        self.scale_var = tk.StringVar(value="None")
        ttk.Label(customize_train_set_frame, text="Scale Type").grid(column=0, row=2)
        ttk.OptionMenu(customize_train_set_frame, self.scale_var, "None", "None","StandardScaler", "MinMaxScaler").grid(column=1, row=2)

        self.difference_choice_var = tk.IntVar(value=0)
        self.interval_var = tk.IntVar(value="")
        tk.Checkbutton(customize_train_set_frame, text='Difference', variable=self.difference_choice_var, offvalue=0, onvalue=1, command=self.openDifference).grid(column=0, row=3)
        ttk.Label(customize_train_set_frame, text="Interval").grid(column=1, row=3)
        self.interval_entry = ttk.Entry(customize_train_set_frame, textvariable=self.interval_var, state=tk.DISABLED)
        self.interval_entry.grid(column=2, row=3)
        

        # Lag Options
        lag_options_frame = ttk.Labelframe(self.root, text="Lag Options")
        lag_options_frame.grid(column=0, row=2)

        self.acf_lags = tk.IntVar(value="")
        ttk.Label(lag_options_frame, text="Number of Lags").grid(column=0, row=0)
        ttk.Entry(lag_options_frame, textvariable=self.acf_lags).grid(column=1, row=0)
        ttk.Button(lag_options_frame, text="Show ACF", command=lambda: self.showACF(self.acf_lags.get())).grid(column=2, row=0)

        self.lag_option_var = tk.IntVar(value="")
        tk.Radiobutton(lag_options_frame, text="Use All Lags", value=0, variable=self.lag_option_var, command=self.openEntries).grid(column=0, row=1)
        tk.Radiobutton(lag_options_frame, text="Use Selected(1,3,..)", value=1, variable=self.lag_option_var, command=self.openEntries).grid(column=0, row=2)
        tk.Radiobutton(lag_options_frame, text="Use Best N", value=2, variable=self.lag_option_var, command=self.openEntries).grid(column=0, row=3)
        tk.Radiobutton(lag_options_frame, text="Use Correlation > n", value=3, variable=self.lag_option_var, command=self.openEntries).grid(column=0, row=4)
        
        self.lag_entries = [ttk.Entry(lag_options_frame, state=tk.DISABLED) for i in range(3)]
        [self.lag_entries[i-2].grid(column=1, row=i) for i in range(2,5)]

        # Create Model
        create_model_frame = ttk.Labelframe(self.root, text="Create Model")
        create_model_frame.grid(column=1, row=0)
        
        self.model_instance = 0
        self.runtime = datetime.now().strftime("%d/%m/%Y %H:%M")
        self.do_optimization = False
        
        ## Model Without Optimization
        model_without_optimization_frame = ttk.Labelframe(create_model_frame, text="Model Without Optimization")
        model_without_optimization_frame.grid(column=0, row=0)

        ttk.Label(model_without_optimization_frame, text="Number of Hidden Layer").grid(column=0, row=0)
        
        no_optimization_names = {1:"Neurons in First Layer", 2:"Neurons in Second Layer", 3:"Neurons in Third Layer"}

        self.neuron_numbers_var = [tk.IntVar(value="") for i in range(3)]
        self.activation_var = [tk.StringVar(value="relu") for i in range(3)]
        self.no_optimization_choice_var = tk.IntVar(value=0)
        
        self.no_optimization = [
                [
                    tk.Radiobutton(model_without_optimization_frame, text=i+1, value=i+1, variable=self.no_optimization_choice_var, command=lambda: self.openOptimizationLayers(True)).grid(column=i+1, row=0),
                    ttk.Label(model_without_optimization_frame, text=no_optimization_names[i+1]).grid(column=0, row=i+1),
                    ttk.Entry(model_without_optimization_frame, textvariable=self.neuron_numbers_var[i], state=tk.DISABLED),
                    ttk.Label(model_without_optimization_frame, text="Activation Function").grid(column=2, row=i+1),
                    ttk.OptionMenu(model_without_optimization_frame, self.activation_var[i], "relu", "relu", "tanh", "sigmoid").grid(column=3, row=i+1)
                ] for i in range(3)
        ]

        for i,j in enumerate(self.no_optimization):
            j[2].grid(column=1, row=i+1)


        ## Model With Optimization
        model_with_optimization_frame = ttk.Labelframe(create_model_frame, text="Model With Optimization")
        model_with_optimization_frame.grid(column=0, row=1)

        optimization_names = {1:"One Hidden Layer", 2:"Two Hidden Layer", 3:"Three Hidden Layer"}
        self.optimization_choice_var = tk.IntVar(value=0)

        self.neuron_min_number_var = [tk.IntVar(value="") for i in range(3)]
        self.neuron_max_number_var = [tk.IntVar(value="") for i in range(3)]

        self.optimization = [
                [
                    tk.Radiobutton(model_with_optimization_frame, text=optimization_names[i+1], value=i+1, variable=self.optimization_choice_var, command=lambda: self.openOptimizationLayers(False)).grid(column=i*2+1, row=0),
                    ttk.Label(model_with_optimization_frame, text=f"N{i+1}_Min").grid(column=i*2, row=1),
                    ttk.Label(model_with_optimization_frame, text=f"N{i+1}_Max").grid(column=i*2, row=2),
                    ttk.Entry(model_with_optimization_frame, textvariable=self.neuron_min_number_var[i], state=tk.DISABLED),
                    ttk.Entry(model_with_optimization_frame, textvariable=self.neuron_max_number_var[i], state=tk.DISABLED)
                ] for i in range(3)
        ]
        
        for i,j in enumerate(self.optimization):
            j[3].grid(column=i*2+1, row=1)
            j[4].grid(column=i*2+1, row=2)

        
        # Hyperparameters
        hyperparameter_frame = ttk.Labelframe(self.root, text="Hyperparameters")
        hyperparameter_frame.grid(column=1, row=1)


        self.hyperparameters = {"Epoch": tk.IntVar(), "Batch_Size": tk.IntVar(), "Optimizer": tk.StringVar(), "Loss_Function": tk.StringVar(), "Learning_Rate": tk.Variable(value=0.001), "Momentum": tk.Variable(value=0.0)}
        
        ttk.Label(hyperparameter_frame, text="Epoch").grid(column=0, row=0)
        ttk.Entry(hyperparameter_frame, textvariable=self.hyperparameters["Epoch"]).grid(column=1, row=0)

        ttk.Label(hyperparameter_frame, text="Batch Size").grid(column=2, row=0)
        ttk.Entry(hyperparameter_frame, textvariable=self.hyperparameters["Batch_Size"]).grid(column=3, row=0)

        ttk.Label(hyperparameter_frame, text="Optimizer").grid(column=0, row=1)
        ttk.OptionMenu(hyperparameter_frame, self.hyperparameters["Optimizer"], "Adam", "Adam", "SGD", "RMSprop").grid(column=1, row=1)

        ttk.Label(hyperparameter_frame, text="Loss_Function").grid(column=2, row=1)
        ttk.OptionMenu(hyperparameter_frame, self.hyperparameters["Loss_Function"], "mean_squared_error", "mean_squared_error", "mean_absolute_error", "mean_absolute_percentage_error").grid(column=3, row=1)

        ttk.Label(hyperparameter_frame, text="Learning Rate").grid(column=0, row=2)
        ttk.Entry(hyperparameter_frame, textvariable=self.hyperparameters["Learning_Rate"]).grid(column=1, row=2)

        ttk.Label(hyperparameter_frame, text="Momentum").grid(column=2, row=2)
        ttk.Entry(hyperparameter_frame, textvariable=self.hyperparameters["Momentum"]).grid(column=3, row=2)

        model_names = ["MLP Model", "CNN Model", "LSTM Model", "Bi-LSTM Model"]
        self.model_var = tk.IntVar(value="")
        ttk.Label(hyperparameter_frame, text="Model Type").grid(column=0, row=3, columnspan=4)
        [tk.Radiobutton(hyperparameter_frame, text=model_names[i], value=i, variable=self.model_var).grid(column=i, row=4) for i in range(4)]

        self.train_loss = tk.Variable(value="")
        ttk.Button(hyperparameter_frame, text="Create Model", command=self.createModel).grid(column=0, row=5)
        ttk.Label(hyperparameter_frame, text="Train Loss").grid(column=1, row=5)
        ttk.Entry(hyperparameter_frame, textvariable=self.train_loss).grid(column=2, row=5)

        ttk.Label(hyperparameter_frame, text="Best Model Neuron Numbers").grid(column=0, row=6)
        self.best_model_neurons = [tk.IntVar(value="") for i in range(3)]
        [ttk.Entry(hyperparameter_frame, textvariable=self.best_model_neurons[i], width=5).grid(column=i+1, row=6) for i in range(3)]
       
        # Test Model
        test_model_frame = ttk.Labelframe(self.root, text="Test Model")
        test_model_frame.grid(column=1, row=2)
        
        forecast_num = tk.IntVar(value="")
        ttk.Label(test_model_frame, text="# of Forecast").grid(column=0, row=0)
        ttk.Entry(test_model_frame, textvariable=forecast_num).grid(column=1, row=0)
        ttk.Button(test_model_frame, text="Values", command=self.showTestSet).grid(column=2, row=0)

        test_file_path = tk.StringVar()
        ttk.Label(test_model_frame, text="Test File Path").grid(column=0, row=1)
        ttk.Entry(test_model_frame, textvariable=test_file_path).grid(column=1, row=1)
        ttk.Button(test_model_frame, text="Get Test Set", command=lambda: self.getTestSet(test_file_path)).grid(column=2, row=1)

        ttk.Button(test_model_frame, text="Test Model", command=lambda: self.testModel(forecast_num.get())).grid(column=0, row=2)
        ttk.Button(test_model_frame, text="Actual vs Forecasted Graph", command=self.vsGraph).grid(column=1, row=2)

        self.mape = tk.Variable(value="")
        ttk.Label(test_model_frame, text="Test MAPE").grid(column=0, row=3)
        ttk.Entry(test_model_frame, textvariable=self.mape).grid(column=1, row=3)
        ttk.Button(test_model_frame, text="My Func", command=self.myFunc).grid(column=0, row=4)



    def myFunc(self):
        pred_name = "TRAFLOAD_CELL_PS_DL_KB"
        m_name = "Bi-LSTM"
        self.model_var.set(3)
        count = 1
        for i in rows:
            k = i.cells
            percent = int(k[1].text[:-1])
            lag = int(k[4].text)
            layer_num = int(k[6].text)
            neurons = k[7].text
            af = ""
            epoch = int(k[9].text)
            batch_size = int(k[10].text)
            loss_function = k[12].text.replace("\n", " ").strip()
            pred_num = int(k[14].text)
            
            try:
                interval = int(k[3].text)
                self.difference_choice_var.set(1)
                self.interval_var.set(interval)
            except:
                interval = "-"
                self.difference_choice_var.set(0)

            self.train_size_var.set(percent)
            self.acf_lags.set(lag)
            self.lag_option_var.set(0)
            self.no_optimization_choice_var.set(layer_num)
            self.openOptimizationLayers(True)
            n_nums = [int(j) for j in k[7].text.split(',')]
            for j in range(layer_num):
                self.neuron_numbers_var[j].set(n_nums[j])
                af += "Relu,"

            self.hyperparameters["Epoch"].set(epoch)
            self.hyperparameters["Batch_Size"].set(batch_size)
            lss = {"Mean Squared error":"mean_squared_error", "Mean Absolute error":"mean_absolute_error", "MAPE":"mean_absolute_percentage_error"} 
            self.hyperparameters["Loss_Function"].set(lss[loss_function])
            self.createModel()
            self.testModel(pred_num)
            print(count, ". finished")
            print("Test MAPE =", self.mape.get())
            count += 1

            st = savetable.add_row().cells
            st[0].text = pred_name
            st[1].text = str(percent) + "%"
            st[2].text = "-"
            st[3].text = str(interval)
            st[4].text = str(lag)
            st[5].text = "Use All Lags"
            st[6].text = str(layer_num)
            st[7].text = neurons
            st[8].text = str(af[:-1])
            st[9].text = str(epoch)
            st[10].text = str(batch_size)
            st[11].text = "Adam"
            st[12].text = loss_function
            st[13].text = "0.001"
            st[14].text = str(pred_num)
            st[15].text = str(self.mape.get())[:4]
            
            clear_session()

        savedoc.save("Real Data/Forecast_Bidirectional LSTM.docx")

    def readCsv(self, file_path):
        path = filedialog.askopenfilename(filetypes=[("Csv Files", "*.csv")])
        file_path.set(path)
        self.df = pd.read_csv(path)
        
        self.input_list.delete(0, tk.END)

        for i in self.df.columns:
            self.input_list.insert(tk.END, i)

    def getTestSet(self, file_path):
        path = filedialog.askopenfilename(filetypes=[("Csv Files", "*.csv")])
        file_path.set(path)
        self.test_df = pd.read_csv(path)

    def showTestSet(self):
        top = tk.Toplevel(self.root)
        df = pd.DataFrame({"Test": self.forecast[:,0], "Predict": self.pred[:,0]})
        pt = Table(top, dataframe=df, editable=False)
        pt.show()

    def addPredictor(self):
        try:
            a = self.input_list.get(self.input_list.curselection())
            if a not in self.predictor_list.get(0,tk.END):
                self.predictor_list.insert(tk.END, a)
        except:
            pass

    def ejectPredictor(self):
        try:
            self.predictor_list.delete(self.predictor_list.curselection())
        except:
            pass
    
    def addTarget(self):
        try:
            a = self.input_list.get(self.input_list.curselection())
            if self.target_list.size() < 1:
                self.target_list.insert(tk.END, a)
        except:
            pass

    def ejectTarget(self):
        try:
            self.target_list.delete(self.target_list.curselection())
        except:
            pass

    def openDifference(self):
        s = tk.NORMAL if self.difference_choice_var.get() else tk.DISABLED
        self.interval_entry["state"] = s

    def showACF(self, lags):
        top = tk.Toplevel()
        fig = plt.Figure((20, 15))
        
        data = self.df[self.target_list.get(0)]

        ax = fig.add_subplot(211)
        plot_acf(data, ax=ax, lags=lags)

        ax1 = fig.add_subplot(212)
        plot_pacf(data, ax=ax1, lags=lags)

        canvas = FigureCanvasTkAgg(fig, top)
        canvas.draw()
        canvas.get_tk_widget().pack(side=tk.TOP, fill=tk.BOTH, expand=True)
        toolbar = NavigationToolbar2Tk(canvas, top)
        toolbar.update()
        canvas._tkcanvas.pack(side=tk.TOP, fill=tk.BOTH, expand=True)

    def openEntries(self):
        o = self.lag_option_var.get() - 1
        for i, j in enumerate(self.lag_entries):
            if i == o:
                j["state"] = tk.NORMAL
            else:
                j["state"] = tk.DISABLED

    def openOptimizationLayers(self, var):
        for i in self.no_optimization:
            i[2]["state"] = tk.DISABLED

        for i in self.optimization:
            i[3]["state"] = tk.DISABLED
            i[4]["state"] = tk.DISABLED

        if var:
            for i in range(self.no_optimization_choice_var.get()):
                self.no_optimization[i][2]["state"] = tk.NORMAL
            self.optimization_choice_var.set(0)
            
            self.do_optimization = False

        if not var:
            for i in range(self.optimization_choice_var.get()):
                self.optimization[i][3]["state"] = tk.NORMAL
                self.optimization[i][4]["state"] = tk.NORMAL
            self.no_optimization_choice_var.set(0)
            self.do_optimization = True

    def difference(self, data, diff, fill_values=None):
        interval = self.interval_var.get()
        if diff:
            return np.array([data[i] - data[i-interval] for i in range(interval, len(data))])
        else:
            for i in range(len(data)):
                if i >= interval:
                    data[i] = data[i] + data[i-interval]
                else:
                    data[i] = data[i] + fill_values[(len(fill_values) - interval)+i]


    def getDataset(self):
        choice = self.scale_var.get()
        difference_choice = self.difference_choice_var.get()

        size_choice = self.size_choice_var.get()
        size = self.train_size_var.get() if size_choice else (self.train_size_var.get()/100) * len(self.df)
        size = int(size)

        placeholder = [i for i in self.predictor_list.get(0, tk.END)]
        features = self.df[placeholder].iloc[:size].to_numpy()
        label = self.df[[self.target_list.get(0)]].iloc[:size].to_numpy()
        
        self.fill_values = label

        if difference_choice:
            features = self.difference(features, True)
            label = self.difference(label, True)

        if choice == "StandardScaler":
            self.feature_scaler = StandardScaler()
            self.label_scaler = StandardScaler()
            
            features = self.feature_scaler.fit_transform(features)
            label = self.label_scaler.fit_transform(label)
        
        elif choice == "MinMaxScaler":
            self.feature_scaler = MinMaxScaler()
            self.label_scaler = MinMaxScaler()
            
            features = self.feature_scaler.fit_transform(features)
            label = self.label_scaler.fit_transform(label)

        print(len(features), len(label))
        return features, label

    def createLag(self, features, label):
        X, y = [], []
        n = self.acf_lags.get()
        for i in range(len(features) - n):
            X.append(features[i:i+n])
            y.append(label[i+n])
        
        self.last = np.array(features[len(features)-n:])

        X, y = np.array(X), np.array(y)

        lag_type = self.lag_option_var.get()
        acf_vals = acf(self.df[self.target_list.get(0)].values, nlags=n, fft=False)

        if lag_type == 1:
            lag = self.lag_entries[0].get()
            numbers = [int(i) for i in lag.split(',')]
            X = X[:, numbers, :]

        elif lag_type == 2:
            lag = self.lag_entries[1].get()
            numbers = np.argsort(acf_vals[1:])[-int(lag):]
            X = X[:, numbers, :]

        elif lag_type == 3:
            lag = self.lag_entries[2].get()
            numbers = [i for i,j in enumerate(acf_vals[1:]) if j > float(lag)]
            X = X[:, numbers, :]

        print(X.shape, y.shape)
        
        return X, y

    def createModel(self):
        self.model_instance += 1

        features, label = self.getDataset()
        X_train, y_train = self.createLag(features, label)
        self.X_train = X_train

        learning_rate = float(self.hyperparameters["Learning_Rate"].get())
        momentum = float(self.hyperparameters["Momentum"].get())
        optimizers = {
                "Adam": Adam(learning_rate=learning_rate),
                "SGD": SGD(learning_rate=learning_rate, momentum=momentum),
                "RMSprop": RMSprop(learning_rate=learning_rate, momentum=momentum)
                }

        shape = (X_train.shape[1], X_train.shape[2])
        model_choice = self.model_var.get()

        if not self.do_optimization:
            model = Sequential()
            model.add(Input(shape=shape))
            
            if model_choice == 0:
                model.add(Flatten())

            layers = self.no_optimization_choice_var.get()
            for i in range(layers):
                neuron_number = self.neuron_numbers_var[i].get()
                activation_function = self.activation_var[i].get()
                if model_choice == 0:
                    model.add(Dense(neuron_number, activation=activation_function))
                
                elif model_choice == 1:
                    model.add(Conv1D(filters=neuron_number, kernel_size=2, activation=activation_function))
                    model.add(MaxPooling1D(pool_size=2))
                
                elif model_choice == 2:
                    if i == layers-1:
                        model.add(LSTM(neuron_number, activation=activation_function, return_sequences=False))
                    else:
                        model.add(LSTM(neuron_number, activation=activation_function, return_sequences=True))

                elif model_choice == 3:
                    if i == layers-1:
                        model.add(Bidirectional(LSTM(neuron_number, activation=activation_function, return_sequences=False)))
                    else:
                        model.add(Bidirectional(LSTM(neuron_number, activation=activation_function, return_sequences=True)))
            
            if model_choice == 1:
                model.add(Flatten())
                model.add(Dense(32))

            model.add(Dense(1))
            model.compile(optimizer = optimizers[self.hyperparameters["Optimizer"].get()], loss=self.hyperparameters["Loss_Function"].get())
            
            history = model.fit(X_train, y_train, epochs=self.hyperparameters["Epoch"].get(), batch_size=self.hyperparameters["Batch_Size"].get(), verbose=1)
            loss = history.history["loss"][-1]
            self.train_loss.set(loss)

        elif self.do_optimization:
            layer = self.optimization_choice_var.get()

            if model_choice == 0:
                def build_model(hp):
                    model = Sequential()
                    model.add(Input(shape=shape))
                    model.add(Flatten())
                    for i in range(layer):
                        n_min = self.neuron_min_number_var[i].get()
                        n_max = self.neuron_max_number_var[i].get()
                        step = int((n_max - n_min)/4)
                        model.add(Dense(units=hp.Int('MLP_'+str(i), min_value=n_min, max_value=n_max, step=step), activation='relu'))
                    model.add(Dense(1))
                    model.compile(optimizer = optimizers[self.hyperparameters["Optimizer"].get()], loss=self.hyperparameters["Loss_Function"].get())
                    return model
                
                name = str(self.model_instance) + ". MLP"

            elif model_choice == 1:
                def build_model(hp):
                    model = Sequential()
                    model.add(Input(shape=shape))
                    for i in range(layer):
                        n_min = self.neuron_min_number_var[i].get()
                        n_max = self.neuron_max_number_var[i].get()
                        step = int((n_max-n_min)/4)
                        model.add(Conv1D(filters=hp.Int("CNN_"+str(i), min_value=n_min, max_value=n_max, step=step), kernel_size=2, activation="relu"))
                        model.add(MaxPooling1D(pool_size=2))
                    
                    model.add(Flatten())
                    model.add(Dense(32))
                    model.add(Dense(1))
                    model.compile(optimizer = optimizers[self.hyperparameters["Optimizer"].get()], loss=self.hyperparameters["Loss_Function"].get())
                    return model
                
                name = str(self.model_instance) + ". CNN"

            elif model_choice == 2:
                def build_model(hp):
                    model = Sequential()
                    model.add(Input(shape=shape))
                    for i in range(layer):
                        n_min = self.neuron_min_number_var[i].get()
                        n_max = self.neuron_max_number_var[i].get()
                        step = int((n_max - n_min)/4)
                        model.add(LSTM(units=hp.Int("LSTM_"+str(i), min_value=n_min, max_value=n_max, step=step), activation='relu', return_sequences=True))
                        if i == layer-1:
                            model.add(LSTM(units=hp.Int("LSTM_"+str(i), min_value=n_min, max_value=n_max, step=step), activation='relu', return_sequences=False))
                    
                    model.add(Dense(1))
                    model.compile(optimizer = optimizers[self.hyperparameters["Optimizer"].get()], loss=self.hyperparameters["Loss_Function"].get())
                    return model
                
                name = str(self.model_instance) + ". LSTM"
            
            elif model_choice == 3:
                def build_model(hp):
                    model = Sequential()
                    model.add(Input(shape=shape))
                    for i in range(layer):
                        n_min = self.neuron_min_number_var[i].get()
                        n_max = self.neuron_max_number_var[i].get()
                        step = int((n_max - n_min)/4)
                        model.add(Bidirectional(LSTM(units=hp.Int("LSTM_"+str(i), min_value=n_min, max_value=n_max, step=step), activation='relu', return_sequences=True)))
                        if i == layer-1:
                            model.add(Bidirectional(LSTM(units=hp.Int("LSTM_"+str(i), min_value=n_min, max_value=n_max, step=step), activation='relu', return_sequences=False)))
                    
                    model.add(Dense(1))
                    model.compile(optimizer = optimizers[self.hyperparameters["Optimizer"].get()], loss=self.hyperparameters["Loss_Function"].get())
                    return model

                name = str(self.model_instance) + ". Bi-LSTM"


            tuner = RandomSearch(build_model, objective='loss', max_trials=3, executions_per_trial=1, directory=self.runtime, project_name=name)
            
            tuner.search(X_train, y_train, epochs=self.hyperparameters["Epoch"].get(), batch_size=self.hyperparameters["Batch_Size"].get())
            hps = tuner.get_best_hyperparameters(num_trials = 1)[0]
            model = tuner.hypermodel.build(hps)
            
            history = model.fit(X_train, y_train, epochs=self.hyperparameters["Epoch"].get(), batch_size=self.hyperparameters["Batch_Size"].get(), verbose=1)
            loss = history.history["loss"][-1]
            self.train_loss.set(loss)
            

            for i in range(layer):
                if model_choice == 0:
                    self.best_model_neurons[i].set(model.get_layer(index=i+1).get_config()["units"])
                elif model_choice == 1:
                    self.best_model_neurons[i].set(model.get_layer(index=(2*i)).get_config()["filters"])
                elif model_choice == 2:
                    self.best_model_neurons[i].set(model.get_layer(index=i).get_config()["units"])
                elif model_choice == 3:
                    self.best_model_neurons[i].set(model.get_layer(index=i).get_config()["layer"]["config"]["units"])

        self.model = model

    def testModel(self, num):
        input_value = self.last
        steps, features = self.X_train.shape[1], self.X_train.shape[2]
        shape = (1,steps,features)
        pred = []

        for _ in range(num):
            output = self.model.predict(input_value.reshape(shape), verbose=0)
            pred = np.append(pred, output)
            input_value = np.append(input_value, output)[-shape[1]:]

        self.pred = np.array(pred).reshape(-1,1)
        
        if self.scale_var.get() != "None":
            self.pred = self.label_scaler.inverse_transform(self.pred)
        
        if self.difference_choice_var.get():
            self.difference(self.pred, False, self.fill_values)
        
        self.forecast = self.test_df[[self.target_list.get(0)]]
        self.forecast = np.asarray(self.forecast)[:num]

        print(self.pred.shape, self.forecast.shape)

        mape = np.mean(np.abs((self.forecast - self.pred) / self.forecast)) * 100
        self.mape.set(mape)

    def vsGraph(self):
        
        plt.plot(self.forecast)
        plt.plot(self.pred)
        plt.legend(["Test", "Predict"], loc="upper left")
        plt.show()

class SupportVectorMachine:
    def __init__(self):
        self.root = ttk.Frame()
        
        # Get Train Set
        get_train_set_frame = ttk.Labelframe(self.root, text="Get Train Set")
        get_train_set_frame.grid(column=0, row=0)

        file_path = tk.StringVar(value="")
        ttk.Label(get_train_set_frame, text="Train File Path").grid(column=0, row=0)
        ttk.Entry(get_train_set_frame, textvariable=file_path).grid(column=1, row=0)
        ttk.Button(get_train_set_frame, text="Read Csv", command=lambda: self.readCsv(file_path)).grid(column=2, row=0)

        self.input_list = tk.Listbox(get_train_set_frame)
        self.input_list.grid(column=0, row=1)

        self.predictor_list = tk.Listbox(get_train_set_frame)
        self.predictor_list.grid(column=1, row=1)

        self.target_list = tk.Listbox(get_train_set_frame)
        self.target_list.grid(column=2, row=1)

        ttk.Button(get_train_set_frame, text="Add Predictor", command=self.addPredictor).grid(column=1, row=2)
        ttk.Button(get_train_set_frame, text="Eject Predictor", command=self.ejectPredictor).grid(column=1, row=3)

        ttk.Button(get_train_set_frame, text="Add Target", command=self.addTarget).grid(column=2, row=2)
        ttk.Button(get_train_set_frame, text="Eject Target", command=self.ejectTarget).grid(column=2, row=3)
       
        # Customize Train Set
        customize_train_set_frame = ttk.Labelframe(self.root, text="Customize Train Set")
        customize_train_set_frame.grid(column=0, row=1)

        self.train_size_var = tk.IntVar(value="")
        ttk.Label(customize_train_set_frame, text="# of Rows in Train Set").grid(column=0, row=0)
        ttk.Entry(customize_train_set_frame, textvariable=self.train_size_var).grid(column=1, row=0)

        self.size_choice_var = tk.IntVar(value=0)
        tk.Radiobutton(customize_train_set_frame, text="As Percent", value=0, variable=self.size_choice_var).grid(column=0, row=1)
        tk.Radiobutton(customize_train_set_frame, text="As Number", value=1, variable=self.size_choice_var).grid(column=1, row=1)

        self.scale_var = tk.StringVar(value="None")
        ttk.Label(customize_train_set_frame, text="Scale Type").grid(column=0, row=2)
        ttk.OptionMenu(customize_train_set_frame, self.scale_var, "None", "None","StandardScaler", "MinMaxScaler").grid(column=1, row=2)

        # Model
        self.model_frame = ttk.Labelframe(self.root, text="Model Frame")
        self.model_frame.grid(column=1, row=0)

        ## Model Type
        self.model_type_frame = ttk.Labelframe(self.model_frame, text="Type of SVR Model")
        self.model_type_frame.grid(column=0, row=0)

        self.model_type_var = tk.IntVar(value=0)
        tk.Radiobutton(self.model_type_frame, text="Epsilon-SVR", value=0, variable=self.model_type_var).grid(column=0, row=0)
        tk.Radiobutton(self.model_type_frame, text="Nu-SVR", value=1, variable=self.model_type_var).grid(column=1, row=0)

        ## Kernel Type
        self.kernel_type_frame = ttk.Labelframe(self.model_frame, text="Kernel Function")
        self.kernel_type_frame.grid(column=0, row=1)

        self.kernel_type_var = tk.IntVar()
        tk.Radiobutton(self.kernel_type_frame, text="Linear", value=0, variable=self.kernel_type_var).grid(column=0, row=0, sticky=tk.W)
        tk.Radiobutton(self.kernel_type_frame, text="RBF", value=1, variable=self.kernel_type_var).grid(column=1, row=0, sticky=tk.W)
        tk.Radiobutton(self.kernel_type_frame, text="Polynomial", value=2, variable=self.kernel_type_var).grid(column=0, row=1)
        tk.Radiobutton(self.kernel_type_frame, text="Sigmoid", value=3, variable=self.kernel_type_var).grid(column=1, row=1)
        
        ## Parameter Optimization
        self.parameter_optimization_frame = ttk.Labelframe(self.model_frame, text="Parameter Optimization")
        self.parameter_optimization_frame.grid(column=1, row=0, rowspan=2)

        self.grid_option_var = tk.IntVar(value=0)
        ttk.Checkbutton(self.parameter_optimization_frame, text="Do grid search for optimal parameters", offvalue=0, onvalue=1, variable=self.grid_option_var).grid(column=0, row=0, columnspan=3)

        self.grid_range_var = tk.IntVar(value="")
        self.grid_times_var = tk.IntVar(value="")
        ttk.Label(self.parameter_optimization_frame, text="Intervals:").grid(column=0, row=1)
        ttk.Entry(self.parameter_optimization_frame, textvariable=self.grid_range_var, width=8, state=tk.DISABLED).grid(column=1, row=1)
        ttk.Entry(self.parameter_optimization_frame, textvariable=self.grid_times_var, width=8, state=tk.DISABLED).grid(column=2, row=1)

    def readCsv(self, file_path):
        path = filedialog.askopenfilename(filetypes=[("Csv Files", "*.csv")])
        file_path.set(path)
        self.df = pd.read_csv(path)
        
        self.input_list.delete(0, tk.END)

        for i in self.df.columns:
            self.input_list.insert(tk.END, i)

    def getTestSet(self, file_path):
        path = filedialog.askopenfilename(filetypes=[("Csv Files", "*.csv")])
        file_path.set(path)
        self.test_df = pd.read_csv(path)

    def showTestSet(self):
        top = tk.Toplevel(self.root)
        df = pd.DataFrame({"Test": self.label[:,0], "Predict": self.pred[:,0]})
        pt = Table(top, dataframe=df, editable=False)
        pt.show()

    def addPredictor(self):
        try:
            a = self.input_list.get(self.input_list.curselection())
            if a not in self.predictor_list.get(0,tk.END):
                self.predictor_list.insert(tk.END, a)
        except:
            pass

    def ejectPredictor(self):
        try:
            self.predictor_list.delete(self.predictor_list.curselection())
        except:
            pass
    
    def addTarget(self):
        try:
            a = self.input_list.get(self.input_list.curselection())
            if self.target_list.size() < 1:
                self.target_list.insert(tk.END, a)
        except:
            pass

    def ejectTarget(self):
        try:
            self.target_list.delete(self.target_list.curselection())
        except:
            pass

 

s = GUI()
s.start()

